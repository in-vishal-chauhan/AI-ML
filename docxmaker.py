from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from docx2pdf import convert

# Initialize document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading("Extendable Tools (Orchestrator)", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_paragraph(
    "This guide outlines the standardized process for extending the AI Orchestrator system with custom tools. "
    "By following the steps below, developers can register new functions that the orchestrator will recognize and "
    "invoke automatically based on user input. The system is designed with extensibility in mind, requiring minimal configuration for new capabilities."
)

# Step 1: Define the Function
doc.add_heading("Step 1: Define the Function", level=2)
doc.add_paragraph(
    "Create a Python function that accepts a single argument (`user_input: str`) and returns a string response. "
    "This function encapsulates the logic you want the orchestrator to perform."
)

doc.add_paragraph(
    '''def calculate_profit(user_input: str) -> str:\n'''
    '''    """Calculates profit based on cost and selling price in the input."""\n'''
    '''    return "Profit is ₹30"''',
    style='Intense Quote'
)

# Step 2: Add Description to system_prompt
doc.add_heading("Step 2: Update the `system_prompt`", level=2)
doc.add_paragraph(
    "Update the orchestrator’s `system_prompt` to include a brief description of your function. "
    "This informs the LLM of the tool’s purpose and when to invoke it."
)

doc.add_paragraph(
    '''system_prompt = """\n'''
    '''Functions:\n'''
    '''- check_in_document: For company or document-related queries.\n'''
    '''- check_in_db: For product rates or inventory information.\n'''
    '''- suggest_clothing_combination: For outfit recommendations.\n'''
    '''- get_current_date: Returns today’s date.\n'''
    '''- calculate_profit: For computing profit based on input.\n'''
    '''- your_new_function_name: Describe the function’s purpose here.\n"""''',
    style='Intense Quote'
)

# Step 3: Register in tool_labels
doc.add_heading("Step 3: Register the Tool in `tool_labels`", level=2)
doc.add_paragraph(
    "Add your new tool to the `tool_labels` dictionary. The key is the function name, and the value is the user-facing label."
)

doc.add_paragraph(
    '''tool_labels = {\n'''
    '''    "check_in_document": "Company Info",\n'''
    '''    "check_in_db": "Product Rates",\n'''
    '''    "suggest_clothing_combination": "Clothing Suggestions",\n'''
    '''    "get_current_date": "Current Date",\n'''
    '''    "calculate_profit": "Profit Calculation",\n'''
    '''    "your_new_function_name": "Your Tool Label",\n'''
    '''}''',
    style='Intense Quote'
)

# Step 4: You're Done
doc.add_heading("Step 4: You're Done!", level=2)
doc.add_paragraph(
    "Once the above steps are complete, the orchestrator will automatically detect and utilize your function "
    "when appropriate based on user intent. No further configuration is needed."
)

# Example Section
doc.add_heading("Example: Adding a Currency Converter Tool", level=2)

doc.add_heading("Function Definition", level=3)
doc.add_paragraph(
    '''def convert_currency(user_input: str) -> str:\n'''
    '''    """Converts currency from INR to USD based on the user input."""\n'''
    '''    return "₹100 = $1.20"''',
    style='Intense Quote'
)

doc.add_heading("Add to `system_prompt`", level=3)
doc.add_paragraph(
    '''- convert_currency: Converts between currencies such as INR, USD, and EUR.''',
    style='Intense Quote'
)

doc.add_heading("Add to `tool_labels`", level=3)
doc.add_paragraph(
    '''tool_labels = {\n'''
    '''    ...\n'''
    '''    "convert_currency": "Currency Converter",\n'''
    '''}''',
    style='Intense Quote'
)

# Summary Table
doc.add_heading("Quick Reference Summary", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Table Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Step"
hdr_cells[1].text = "Action Description"

# Add rows
summary_rows = [
    ("1", "Define a new function with `user_input` as the parameter."),
    ("2", "Add a description of the function to the `system_prompt`."),
    ("3", "Register the function in the `tool_labels` dictionary."),
    ("4", "The orchestrator will now auto-discover and use your tool.")
]

for step, action in summary_rows:
    row_cells = table.add_row().cells
    row_cells[0].text = step
    row_cells[1].text = action

# Footer note
doc.add_paragraph().add_run(
    "For advanced tools or integrations (e.g., APIs, databases), ensure the logic is secure and handles edge cases gracefully."
).italic = True

# Save document
output_path = "./Extendable_Tools_Orchestrator_Guide.docx"
doc.save(output_path)

print(f"✅ Professional guide saved to: {output_path}")
convert("./Extendable_Tools_Orchestrator_Guide.docx")