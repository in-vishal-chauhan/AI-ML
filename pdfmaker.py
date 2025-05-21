from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
import os

# Create a Word Document
doc = Document()

# Title
doc.add_heading("Extendable Tools (Orchestrator)", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

# Intro
doc.add_paragraph(
    "To support future requirements, the orchestrator system is designed to be easily extendable. "
    "Developers can add new tools (functions) with minimal effort, allowing the AI orchestrator to "
    "automatically detect and use them based on user input."
)

# Step 1
doc.add_heading("1. Create a New Function", level=2)
doc.add_paragraph("- The function must accept a single parameter: `user_input`.\n"
                  "- Write the tool’s logic inside the function to generate the required response.")
doc.add_paragraph(
    '''def calculate_profit(user_input: str) -> str:\n    """Calculates profit based on cost and selling price in the input."""\n    return "Profit is ₹30"''',
    style='Intense Quote'
)

# Step 2
doc.add_heading("2. Add a Description in system_prompt", level=2)
doc.add_paragraph("- In the orchestrator’s `system_prompt`, list the function with a short description.\n"
                  "- This helps the LLM (language model) understand when to use the tool.")
doc.add_paragraph(
    '''system_prompt = """\nFunctions:\n- check_in_document: For company or document-related questions.\n- check_in_db: For product prices, rates, or system data.\n- suggest_clothing_combination: For outfit or clothing suggestions.\n- get_current_date: To get the current date.\n- calculate_profit: For profit calculations.\n- your_new_function_name: Describe the purpose here.\n"""''',
    style='Intense Quote'
)

# Step 3
doc.add_heading("3. Register the Tool in tool_labels Dictionary", level=2)
doc.add_paragraph("- Add a key-value pair to the `tool_labels` dictionary.\n"
                  "- The key is the function name.\n"
                  "- The value is the label that will appear in the user-facing response.")
doc.add_paragraph(
    '''tool_labels = {\n    "check_in_document": "Company Info",\n    "check_in_db": "Product Rates",\n    "suggest_clothing_combination": "Clothing Suggestions",\n    "get_current_date": "Current Date",\n    "calculate_profit": "Profit Calculation",\n    "your_new_function_name": "Your Tool Label",\n}''',
    style='Intense Quote'
)

# Step 4
doc.add_heading("4. That’s It!", level=2)
doc.add_paragraph("- The orchestrator will now automatically detect and use the tool when relevant.\n"
                  "- No additional changes are needed in the main orchestration logic.")

# Example
doc.add_heading("Example: Adding a Currency Converter Tool", level=2)

doc.add_heading("Step 1: Function", level=3)
doc.add_paragraph(
    '''def convert_currency(user_input: str) -> str:\n    """Converts currency based on user input."""\n    return "₹100 = $1.20"''',
    style='Intense Quote'
)

doc.add_heading("Step 2: Add to system_prompt", level=3)
doc.add_paragraph(
    '''- convert_currency: For converting between currencies like INR to USD or EUR.''',
    style='Intense Quote'
)

doc.add_heading("Step 3: Add to tool_labels", level=3)
doc.add_paragraph(
    '''tool_labels = {\n    ...\n    "convert_currency": "Currency Converter",\n}''',
    style='Intense Quote'
)

# Summary Table
doc.add_heading("Summary Table", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Step'
hdr_cells[1].text = 'Description'

summary_data = [
    ("1", "Create a function with `user_input` as the parameter."),
    ("2", "Add a description to `system_prompt`."),
    ("3", "Register the tool in `tool_labels`."),
    ("4", "Done! The orchestrator can now use the new tool automatically.")
]

for step, desc in summary_data:
    row_cells = table.add_row().cells
    row_cells[0].text = step
    row_cells[1].text = desc

# Save DOCX and convert to PDF
docx_path = "./Extendable_Tools_Orchestrator_Guide.docx"
doc.save(docx_path)
docx_path
